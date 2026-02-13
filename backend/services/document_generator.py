import os
import re
import logging
import subprocess
from datetime import datetime
from docx import Document
from docx.shared import Pt
from config import Config

logger = logging.getLogger(__name__)

TEMPLATE_DIR = os.path.join(os.path.dirname(os.path.dirname(__file__)), "templates")
OUTPUT_DIR = os.path.join(Config.UPLOAD_DIR, "generated")


def _ensure_output_dir():
    os.makedirs(OUTPUT_DIR, exist_ok=True)


def _get_template_path(doc_type: str) -> str:
    template_map = {
        "tilbud": "offer_template.docx",
        "brev": "letter_template.docx",
        "notat": "note_template.docx",
        "omprofilering": "letter_template.docx",
        "svar_paa_brev": "letter_template.docx",
        "serviceavtale": "offer_template.docx",
    }
    return os.path.join(TEMPLATE_DIR, template_map.get(doc_type, "letter_template.docx"))


def _safe_filename(name: str) -> str:
    return "".join(c if c.isalnum() or c in "._- " else "_" for c in name).strip()


def generate_files(doc: dict) -> dict:
    """Generate Word and PDF files from document data. Returns dict of file paths."""
    _ensure_output_dir()

    base_name = _safe_filename(doc["document_name"])
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    file_base = f"{base_name}_{timestamp}"

    # Generate unsigned Word
    word_path = _generate_word(doc, file_base, signed=False)
    word_signed_path = _generate_word(doc, file_base, signed=True)

    # Generate PDFs from Word files
    pdf_path = _convert_to_pdf(word_path)
    pdf_signed_path = _convert_to_pdf(word_signed_path)

    return {
        "word": word_path,
        "word_signed": word_signed_path,
        "pdf": pdf_path,
        "pdf_signed": pdf_signed_path,
    }


def _replace_paragraph_placeholder(paragraph, key, value, bold=None):
    """Replace placeholder in paragraph, handling split runs."""
    full_text = paragraph.text
    if key not in full_text:
        return False

    new_text = full_text.replace(key, value)
    # Clear all runs and set the replacement text on the first run
    for i, run in enumerate(paragraph.runs):
        if i == 0:
            run.text = new_text
            if bold is not None:
                run.bold = bold
        else:
            run.text = ""
    return True


def _add_markdown_paragraph(document, line, insert_before=None):
    """Add a paragraph with inline markdown formatting (bold/italic)."""
    p = document.add_paragraph()

    # Parse inline markdown: **bold**, *italic*
    # Pattern matches **bold**, *italic*, or plain text
    parts = re.split(r'(\*\*.*?\*\*|\*.*?\*)', line)

    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = p.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("*") and part.endswith("*"):
            run = p.add_run(part[1:-1])
            run.italic = True
        else:
            p.add_run(part)

    return p


def _insert_document_text(document, text, placeholder_para=None):
    """Insert formatted markdown text. If placeholder_para is given, inserts at that
    position and removes the placeholder. Otherwise appends to end."""
    parsed = []
    for line in text.split("\n"):
        line = line.strip()
        # Strip trailing backslash (markdown hard break from rich text editors)
        line = line.rstrip("\\").rstrip()
        if not line:
            parsed.append(("empty", ""))
        elif line.startswith("### "):
            parsed.append(("heading3", line[4:]))
        elif line.startswith("## "):
            parsed.append(("heading2", line[3:]))
        elif line.startswith("# "):
            parsed.append(("heading1", line[2:]))
        elif line.startswith("- "):
            parsed.append(("bullet", line[2:]))
        elif line.startswith("---"):
            continue
        else:
            parsed.append(("normal", line))

    if placeholder_para:
        # Insert at placeholder position using XML manipulation
        body = document.element.body
        ref_elem = placeholder_para._element
        insert_after = ref_elem
        for ptype, content in parsed:
            p = _create_content_paragraph(document, ptype, content)
            elem = p._element
            body.remove(elem)
            insert_after.addnext(elem)
            insert_after = elem
        body.remove(ref_elem)
    else:
        for ptype, content in parsed:
            _create_content_paragraph(document, ptype, content)


def _create_content_paragraph(document, ptype, content):
    """Create a paragraph at end of document and return it."""
    if ptype == "empty":
        return document.add_paragraph()
    elif ptype.startswith("heading"):
        p = document.add_paragraph()
        run = p.add_run(content)
        run.bold = True
        run.font.size = Pt(13)
        return p
    elif ptype == "bullet":
        p = document.add_paragraph()
        _add_inline_runs(p, f"  \u2022  {content}")
        return p
    else:
        p = document.add_paragraph()
        _add_inline_runs(p, content)
        return p


def _add_inline_runs(paragraph, text):
    """Parse inline markdown (bold/italic) and add as runs to paragraph."""
    parts = re.split(r'(\*\*.*?\*\*|\*.*?\*)', text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("*") and part.endswith("*"):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        else:
            paragraph.add_run(part)


def _generate_word(doc: dict, file_base: str, signed: bool) -> str:
    template_path = _get_template_path(doc["document_type"])
    document = Document(template_path)

    # Build replacements matching template placeholders
    replacements = _build_replacements(doc)

    # Placeholders that should NOT be bold (address fields etc.)
    not_bold = {
        "{{recipientPerson}}", "{{recipientAddress}}",
        "{{recipientPostalCode}}", "{{recipientCity}}",
        "{{recipientPhone}}", "{{recipientEmail}}",
    }

    # Replace simple placeholders (handles split runs)
    for paragraph in document.paragraphs:
        for key, value in replacements.items():
            bold_override = False if key in not_bold else None
            if _replace_paragraph_placeholder(paragraph, key, value, bold=bold_override):
                if key == "{{documentTitle}}":
                    for run in paragraph.runs:
                        if run.text.strip():
                            run.bold = True
                            run.font.size = Pt(13)
                    paragraph.paragraph_format.space_after = Pt(6)

    # Also replace in tables (some templates use tables for layout)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for key, value in replacements.items():
                        bold_override = False if key in not_bold else None
                        _replace_paragraph_placeholder(paragraph, key, value, bold=bold_override)

    # Find {{documentText}} placeholder paragraph
    text = doc.get("document_text", "")
    placeholder_para = None

    for paragraph in document.paragraphs:
        if "{{documentText}}" in paragraph.text:
            placeholder_para = paragraph
            break

    # Insert document text at placeholder position (or append to end)
    if text:
        _insert_document_text(document, text, placeholder_para)
    elif placeholder_para:
        document.element.body.remove(placeholder_para._element)

    # Add closing signature block
    document.add_paragraph()
    document.add_paragraph("Med vennlig hilsen")
    document.add_paragraph()

    # Add signature image for signed versions
    if signed:
        sig_path = os.path.join(TEMPLATE_DIR, "signature.png")
        if os.path.exists(sig_path):
            try:
                from docx.shared import Cm
                document.add_picture(sig_path, width=Cm(5))
            except Exception:
                logger.warning("Could not insert signature image")

    p_name = document.add_paragraph()
    run = p_name.add_run("Trond Ilbråten")
    run.bold = True
    document.add_paragraph("Daglig leder")
    document.add_paragraph("Kulde- & Varmepumpeteknikk AS")

    suffix = "_signert" if signed else ""
    output_path = os.path.join(OUTPUT_DIR, f"{file_base}{suffix}.docx")
    document.save(output_path)
    return output_path


def _extract_title(document_name: str) -> str:
    """Extract just the subject/title from the full document name.

    'dd.mm.yyyy - Brev til X vedr. Subject' → 'Subject'
    'dd.mm.yyyy - Tilbud på Subject til X' → 'Subject'
    'dd.mm.yyyy - Notat vedr. Subject' → 'Subject'
    """
    # Remove date prefix
    name = re.sub(r'^\d{2}\.\d{2}\.\d{4}\s*-\s*', '', document_name)
    # "... vedr. Subject"
    if ' vedr. ' in name:
        return name.split(' vedr. ', 1)[1]
    # "Tilbud på Subject til X"
    match = re.match(r'Tilbud på (.+?)(?:\s+til\s+.+)?$', name)
    if match:
        return match.group(1)
    # "Serviceavtale Subject for X"
    match = re.match(r'Serviceavtale (.+?)(?:\s+for\s+.+)?$', name)
    if match:
        return match.group(1)
    # "Svar til X vedr." already handled above, fallback
    return name


def _build_replacements(doc: dict) -> dict:
    """Build replacement dict matching actual template placeholders."""
    today = datetime.now().strftime("%d.%m.%Y")
    doc_name = doc.get("document_name") or ""
    title = _extract_title(doc_name)
    return {
        # Match template placeholders exactly
        "{{dato}}": today,
        "{{mottaker}}": doc.get("recipient_name") or "",
        "{{recipientPerson}}": doc.get("recipient_person") or "",
        "{{recipientAddress}}": doc.get("recipient_address") or "",
        "{{recipientPostalCode}}": doc.get("recipient_postal_code") or "",
        "{{recipientCity}}": doc.get("recipient_city") or "",
        "{{recipientPhone}}": doc.get("recipient_phone") or "",
        "{{recipientEmail}}": doc.get("recipient_email") or "",
        "{{documentTitle}}": title,
        "{{documentName}}": doc_name,
        "{{prisProdukt}}": f"{doc['price_product']:,.2f} kr" if doc.get("price_product") else "",
        "{{prisInstallasjon}}": f"{doc['price_installation']:,.2f} kr" if doc.get("price_installation") else "",
    }


def _convert_to_pdf(word_path: str) -> str | None:
    """Convert Word file to PDF using LibreOffice headless."""
    try:
        result = subprocess.run(
            [
                "libreoffice", "--headless", "--convert-to", "pdf",
                "--outdir", OUTPUT_DIR, word_path,
            ],
            capture_output=True, text=True, timeout=60,
        )
        if result.returncode == 0:
            pdf_path = word_path.rsplit(".", 1)[0] + ".pdf"
            if os.path.exists(pdf_path):
                return pdf_path
        logger.error("LibreOffice conversion failed: %s", result.stderr)
    except subprocess.TimeoutExpired:
        logger.error("LibreOffice conversion timed out for %s", word_path)
    except FileNotFoundError:
        logger.error("LibreOffice not installed - PDF generation unavailable")

    return None
